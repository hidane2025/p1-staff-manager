"""契約書テンプレ取込 & 仮版透かし E2E テスト（DB接続不要）

検証項目:
  1. python-docx で生成した Word → parse_docx → Markdown 化成功
  2. PyPDF2 でテキスト付き PDF → parse_pdf → Markdown 化成功
  3. .md / .txt のデコード
  4. UnsupportedFormatError / 空PDFの例外
  5. 仮版 PDF（is_provisional=True）生成 → 透かしあり
  6. 正規版 PDF（is_provisional=False）生成 → 透かしなし
  7. サンプル PDF を test_e2e/ に保存して目視確認可能にする
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from utils import contract_doc_parser
from utils.contract_doc_parser import (
    DocParseError,
    UnsupportedFormatError,
    parse_docx,
    parse_pdf,
    parse_plain,
    parse_upload,
)
from utils.contract_pdf import (
    ContractVariables,
    build_contract_no,
    generate_contract_pdf,
    render_template,
    today_jst_ymd,
)


OUT = Path(__file__).resolve().parent


SAMPLE_MARKDOWN = """# 業務委託契約書（正規版サンプル）

{{issuer_name}}（以下「甲」という）と、{{staff_name}}（以下「乙」という）は、以下の通り業務委託契約を締結する。

## 第1条（業務内容）

甲は乙に、下記業務を委託する。
- {{event_name}} における {{role}} 業務
- 大会運営全般の補助業務

## 第2条（業務委託料）

乙への報酬は大会終了後に精算するものとする。

---

{{issue_date}}

**甲**: {{issuer_name}}
**乙**: {{staff_name}}

_上記契約の証として、乙が電子署名を行う。_
"""


# ==========================================================================
# 1) Word 取込テスト
# ==========================================================================
def _build_docx_bytes() -> bytes:
    """python-docx でサンプル契約書 Word を生成"""
    from docx import Document
    doc = Document()
    doc.add_heading("業務委託契約書（経理承認版）", level=1)
    doc.add_paragraph(
        "株式会社パシフィック（以下「甲」という）と、{{staff_name}}（以下「乙」という）は、"
        "以下の通り業務委託契約を締結する。"
    )
    doc.add_heading("第1条（業務内容）", level=2)
    doc.add_paragraph("甲は乙に、下記業務を委託する。")
    doc.add_paragraph("ディーラー業務", style="List Bullet")
    doc.add_paragraph("フロア業務", style="List Bullet")

    doc.add_heading("第2条（守秘義務）", level=2)
    p = doc.add_paragraph()
    r = p.add_run("乙は業務上知り得た秘密を")
    r2 = p.add_run("第三者に漏洩してはならない")
    r2.bold = True
    r3 = p.add_run("。")

    # 表サンプル
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "項目"
    table.rows[0].cells[1].text = "内容"
    table.rows[1].cells[0].text = "報酬"
    table.rows[1].cells[1].text = "大会終了後に精算"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_docx() -> None:
    b = _build_docx_bytes()
    md = parse_docx(b)
    assert "# 業務委託契約書" in md, "H1 見出しが Markdown 化されていない"
    assert "## 第1条" in md, "H2 見出しが Markdown 化されていない"
    assert "- ディーラー業務" in md, "箇条書きが正しく変換されていない"
    assert "**第三者に漏洩してはならない**" in md, "bold run が保持されていない"
    assert "| 項目 | 内容 |" in md, "表が Markdown 化されていない"
    assert "{{staff_name}}" in md, "変数プレースホルダが消えている"
    print(f"  [OK] parse_docx: {len(md):,} chars")


def test_parse_upload_docx() -> None:
    b = _build_docx_bytes()
    result = parse_upload("sample.docx", b)
    assert result.parser == "docx"
    assert "業務委託契約書" in result.markdown
    print("  [OK] parse_upload(.docx) dispatch")


# ==========================================================================
# 2) PDF 取込テスト
# ==========================================================================
def _build_pdf_with_text_bytes() -> bytes:
    """reportlab で最低限テキスト抽出可能な PDF を生成"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("Helvetica", 12)
    c.drawString(72, 770, "Contract Template Official v1.0")
    c.drawString(72, 750, "Section 1: Scope of work")
    c.drawString(72, 730, "The Contractor shall provide services.")
    c.drawString(72, 710, "Section 2: Confidentiality")
    c.drawString(72, 690, "All confidential information shall be protected.")
    c.showPage()
    c.save()
    return buf.getvalue()


def test_parse_pdf_with_text() -> None:
    b = _build_pdf_with_text_bytes()
    md = parse_pdf(b)
    assert "Contract Template" in md
    assert "Confidentiality" in md
    print(f"  [OK] parse_pdf: {len(md):,} chars")


def test_parse_pdf_empty_raises() -> None:
    """空 PDF / 画像のみ PDF はエラーになり、手動貼り付けに誘導される"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.showPage()
    c.save()
    try:
        parse_pdf(buf.getvalue())
    except DocParseError as e:
        assert "手動" in str(e) or "抽出" in str(e)
        print("  [OK] 空PDF は DocParseError（手動貼り付けに誘導）")
        return
    raise AssertionError("空PDFで例外が飛ばなかった")


# ==========================================================================
# 3) 平文（.md / .txt）
# ==========================================================================
def test_parse_plain_md() -> None:
    md = parse_plain(SAMPLE_MARKDOWN.encode("utf-8"))
    assert "業務委託契約書" in md
    assert md.endswith("\n")
    print(f"  [OK] parse_plain utf-8: {len(md):,} chars")


def test_parse_plain_cp932() -> None:
    src = "# 業務委託契約書\n\n本契約は日本語のみ。\n"
    md = parse_plain(src.encode("cp932"))
    assert "業務委託契約書" in md
    print("  [OK] parse_plain cp932 decode")


# ==========================================================================
# 4) 異常系
# ==========================================================================
def test_unsupported_extension() -> None:
    try:
        parse_upload("foo.xlsx", b"dummy")
    except UnsupportedFormatError:
        print("  [OK] UnsupportedFormatError for .xlsx")
        return
    raise AssertionError("xlsx なのに例外が飛ばなかった")


# ==========================================================================
# 5,6) 仮版 / 正規版 PDF 透かしテスト
# ==========================================================================
def _build_sample_pdf(is_provisional: bool, out_name: str) -> bytes:
    vars = ContractVariables(
        staff_name="山田 太郎",
        staff_address="東京都新宿区1-2-3",
        role="Dealer",
        event_name="P1 Kyoto 2026",
        issuer_name="株式会社パシフィック",
        issuer_address="東京都港区XX 1-2-3",
    )
    rendered = render_template(SAMPLE_MARKDOWN, vars)
    pdf = generate_contract_pdf(
        rendered_body=rendered,
        contract_no=build_contract_no(1, 1, today_jst_ymd()),
        issuer_name=vars.issuer_name,
        is_provisional=is_provisional,
    )
    (OUT / out_name).write_bytes(pdf)
    return pdf


def _extract_pdf_content_streams(pdf_bytes: bytes) -> str:
    """PyPDF2 で全ページの content stream を decompress して連結する。

    reportlab は stream を zlib 圧縮するため、生バイトの部分一致検索では
    「仮版」ラベルや色指定を検出できない。PyPDF2 でオブジェクトを辿って
    解凍済みの ASCII 文字列にすれば、描画命令（BT/ET, rg, Tj など）を確認できる。
    """
    import PyPDF2
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    parts: list[str] = []
    for page in reader.pages:
        contents = page.get_contents()
        if contents is None:
            continue
        if isinstance(contents, list):
            for obj in contents:
                parts.append(obj.get_data().decode("latin-1", errors="ignore"))
        else:
            parts.append(contents.get_data().decode("latin-1", errors="ignore"))
    return "\n".join(parts)


def _watermark_color_rg(stream: str) -> int:
    """赤橙 #C8381E の RGB 塗り指定 (200/255, 56/255, 30/255 ≈ 0.784 0.220 0.118) を数える。

    reportlab は小数を `.784314` のように先頭 0 を省略して吐くので、
    `0.78...` と `.78...` 両方を許容する正規表現にする。
    命令は `R G B rg`（fill）または `R G B RG`（stroke）。
    """
    import re
    num = r"0?\.\d+"
    pattern = re.compile(
        rf"{num}\s+{num}\s+{num}\s+(?:rg|RG)",
    )
    # 厳密には色値もチェックしたいが、reportlab の出力は "0.784314 0.219608 0.117647"
    # のように安定しているため、赤橙に限定した検知を別関数で行う。
    all_rgb = pattern.findall(stream)
    # rg 呼び出しの前後文字列から赤橙を特定
    hits = 0
    red_orange_pattern = re.compile(
        rf"(?:0?\.78\d+)\s+(?:0?\.21\d+|0?\.22\d+)\s+(?:0?\.1\d+)\s+(?:rg|RG)"
    )
    hits = len(red_orange_pattern.findall(stream))
    _ = all_rgb  # 未使用警告抑制用
    return hits


def test_provisional_watermark() -> None:
    pdf = _build_sample_pdf(True, "test_contract_provisional.pdf")
    assert len(pdf) > 3000
    stream = _extract_pdf_content_streams(pdf)
    # 仮版は透かしボックスに赤橙の塗りを使うので、最低 1 回は rg 指定が出る
    hits = _watermark_color_rg(stream)
    assert hits >= 1, f"仮版透かしの赤橙 rg 指定が見当たらない (hits={hits})"
    # 透明度 /ca （graphics state）も確認（PyPDF2 経由で解凍済み）
    assert "/ca " in stream or "gs" in stream, (
        "透明度（/ca, gs）指定の痕跡が見当たらない"
    )
    print(f"  [OK] 仮版PDF生成 test_contract_provisional.pdf "
          f"({len(pdf):,} bytes, rg hits={hits})")


def test_official_no_watermark() -> None:
    pdf_prov = _build_sample_pdf(True, "_tmp_prov.pdf")
    pdf_off = _build_sample_pdf(False, "test_contract_official.pdf")

    prov_stream = _extract_pdf_content_streams(pdf_prov)
    off_stream = _extract_pdf_content_streams(pdf_off)
    prov_hits = _watermark_color_rg(prov_stream)
    off_hits = _watermark_color_rg(off_stream)

    # 仮版は赤橙塗りが必ず現れる。正規版は見出し下線に赤橙を使うので 1 回は出るが、
    # 仮版は透かしボックス分が加わるため必ず prov > off になる。
    assert prov_hits > off_hits, (
        f"仮版の赤橙出現数が正規版より多くない: "
        f"prov={prov_hits}, off={off_hits}"
    )
    # ファイルサイズもサニティチェック（仮版の方が大きい）
    assert len(pdf_off) < len(pdf_prov), (
        f"正規版 PDF のサイズが仮版以上: prov={len(pdf_prov)}, off={len(pdf_off)}"
    )
    print(f"  [OK] 正規版は透かしなし "
          f"(prov rg={prov_hits}, off rg={off_hits}, "
          f"prov size={len(pdf_prov):,}, off size={len(pdf_off):,})")
    (OUT / "_tmp_prov.pdf").unlink(missing_ok=True)


# ==========================================================================
# 7) 取込→PDF 生成の通しフロー
# ==========================================================================
def test_docx_to_pdf_flow() -> None:
    """Word 取込 → parse_docx → 正規版 PDF 生成まで通す"""
    docx_bytes = _build_docx_bytes()
    markdown = parse_docx(docx_bytes)

    # 変数は生き残っているので、そのままレンダリングできる
    vars = ContractVariables(
        staff_name="テスト 太郎",
        role="Dealer",
        event_name="P1 Kyoto 2026",
    )
    rendered = render_template(markdown, vars)
    assert "テスト 太郎" in rendered

    pdf = generate_contract_pdf(
        rendered_body=rendered,
        contract_no=build_contract_no(99, 99, today_jst_ymd()),
        issuer_name="株式会社パシフィック",
        is_provisional=False,
    )
    (OUT / "test_contract_official_from_docx.pdf").write_bytes(pdf)
    assert len(pdf) > 3000
    print(f"  [OK] docx→md→pdf 通しフロー "
          f"test_contract_official_from_docx.pdf ({len(pdf):,} bytes)")


# ==========================================================================
if __name__ == "__main__":
    print("=== 契約書テンプレ取込 E2E ===")
    test_parse_docx()
    test_parse_upload_docx()
    test_parse_pdf_with_text()
    test_parse_pdf_empty_raises()
    test_parse_plain_md()
    test_parse_plain_cp932()
    test_unsupported_extension()
    test_provisional_watermark()
    test_official_no_watermark()
    test_docx_to_pdf_flow()
    print("\n[ALL PASS]")
